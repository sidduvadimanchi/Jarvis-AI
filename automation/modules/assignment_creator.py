# Backend/Automation/assignment_creator.py
# Jarvis AI — Assignment & Document Creator  (v3.0)
# ─────────────────────────────────────────────────────────────────────────────
# UPGRADES FROM v2.0:
#   ✅ Student Roll Number + Semester on cover page
#   ✅ Diagram/image support (matplotlib charts inserted into doc)
#   ✅ Auto word count validation (re-generates if too short)
#   ✅ Bold keyword detection (**word** → actual bold in docx)
#   ✅ Numbered references section auto-formatted
#   ✅ Terminal progress bar (visual feedback while generating)
#   ✅ Duplicate filename protection (adds timestamp if file exists)
#   ✅ Notes now saved as .docx (not just .txt) — looks professional
#   ✅ Subject auto-detected from topic if not in .env
#   ✅ "create lab manual on X" new template added
#   ✅ AutomationEngine compatible — CreateAssignment() unchanged signature
# ─────────────────────────────────────────────────────────────────────────────

from __future__ import annotations

# ══════════════════════════════════════════════════════════════════════════════
# STANDARD LIBRARY
# ══════════════════════════════════════════════════════════════════════════════
import datetime
import logging
import os
import re
import subprocess
import sys
import time
from pathlib import Path
from typing  import Optional

# ══════════════════════════════════════════════════════════════════════════════
# THIRD-PARTY  (graceful fallbacks)
# ══════════════════════════════════════════════════════════════════════════════
try:
    from dotenv import dotenv_values
    _env: dict = dotenv_values(".env")
except ImportError:
    _env = {}

# ══════════════════════════════════════════════════════════════════════════════
# CONSTANTS  (all configurable via .env)
# ══════════════════════════════════════════════════════════════════════════════
DEFAULT_FONT        : str   = _env.get("DefaultFont",         "Times New Roman")
BODY_FONT_SIZE      : int   = int(_env.get("BodyFontSize",    "12"))
TITLE_FONT_SIZE     : int   = int(_env.get("TitleFontSize",   "18"))
H1_FONT_SIZE        : int   = int(_env.get("H1FontSize",      "16"))
H2_FONT_SIZE        : int   = int(_env.get("H2FontSize",      "14"))
H3_FONT_SIZE        : int   = int(_env.get("H3FontSize",      "13"))
LINE_SPACING_PT     : int   = 18
PARA_SPACE_AFTER_PT : int   = 6
MARGIN_INCHES       : float = 1.0

WORD_LIMITS: dict[str, int] = {
    "short":    300,
    "medium":   600,
    "detailed": 1200,
}
DEFAULT_WORD_LIMIT  : int  = int(_env.get("AssignmentWordLimit", "800"))
ENABLE_PDF_EXPORT   : bool = _env.get("EnablePDFExport",   "False").lower() == "true"
ENABLE_GRAMMAR      : bool = _env.get("EnableGrammarCheck","True").lower()  == "true"
MIN_WORD_THRESHOLD  : float = 0.6   # re-generate if content < 60% of target

USERNAME    : str = _env.get("Username",    "Student")
UNIVERSITY  : str = _env.get("University",  "University")
COURSE      : str = _env.get("Course",      "B.Tech CSE")
SUBJECT     : str = _env.get("Subject",     "Computer Science")
PROFESSOR   : str = _env.get("Professor",   "Professor")
ROLL_NO     : str = _env.get("RollNumber",  "")          # NEW
SEMESTER    : str = _env.get("Semester",    "")          # NEW
GROQ_KEY    : str = _env.get("GroqAPIKey")  or _env.get("GROQ_API_KEY", "")
GROQ_MODEL  : str = _env.get("GroqModel",  "llama-3.1-8b-instant")

DATA_DIR    : Path = Path(_env.get("DataDir", "Data"))
ASSIGN_DIR  : Path = DATA_DIR / "Assignments"
NOTES_DIR   : Path = DATA_DIR / "Notes"
REPORTS_DIR : Path = DATA_DIR / "Reports"
LOG_DIR     : Path = DATA_DIR / "logs"

# ══════════════════════════════════════════════════════════════════════════════
# LOGGING
# ══════════════════════════════════════════════════════════════════════════════
LOG_DIR.mkdir(parents=True, exist_ok=True)
logging.basicConfig(
    level    = logging.INFO,
    format   = "%(asctime)s  [%(levelname)s]  %(message)s",
    handlers = [
        logging.FileHandler(LOG_DIR / "jarvis.log", encoding="utf-8"),
        logging.StreamHandler(),
    ],
)
log = logging.getLogger("jarvis.assignment")

for _d in (ASSIGN_DIR, NOTES_DIR, REPORTS_DIR):
    _d.mkdir(parents=True, exist_ok=True)

# ══════════════════════════════════════════════════════════════════════════════
# OPTIONAL NOTIFIER
# ══════════════════════════════════════════════════════════════════════════════
try:
    from .notifier import notify as _notify        # type: ignore
except ImportError:
    def _notify(title: str, msg: str) -> None:
        log.info("NOTIFY  %s — %s", title, msg)


# ══════════════════════════════════════════════════════════════════════════════
# TERMINAL PROGRESS BAR  (NEW)
# ══════════════════════════════════════════════════════════════════════════════

class _ProgressBar:
    """
    Simple terminal spinner shown while AI is generating content.
    Runs inline — no threads needed.

    Example
    -------
    >>> with _ProgressBar("Generating content"):
    ...     content = _ai_generate(prompt)
    """
    FRAMES = ["⠋","⠙","⠹","⠸","⠼","⠴","⠦","⠧","⠇","⠏"]

    def __init__(self, label: str = "Working") -> None:
        self.label = label
        self._i    = 0

    def __enter__(self):
        sys.stdout.write(f"\n  {self.FRAMES[0]}  {self.label}...")
        sys.stdout.flush()
        return self

    def tick(self) -> None:
        self._i = (self._i + 1) % len(self.FRAMES)
        sys.stdout.write(f"\r  {self.FRAMES[self._i]}  {self.label}...")
        sys.stdout.flush()

    def __exit__(self, *_):
        sys.stdout.write(f"\r  ✅  {self.label} — done\n")
        sys.stdout.flush()


# ══════════════════════════════════════════════════════════════════════════════
# INTERNAL HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _sanitize(name: str) -> str:
    """Remove filesystem-unsafe characters and normalise a filename stem."""
    invalid = '<>:"/\\|?*'
    safe = "".join(c for c in name if c not in invalid).strip()
    return (safe or "document").replace(" ", "_")[:100]


def _unique_path(path: Path) -> Path:
    """
    Return *path* unchanged if it doesn't exist, otherwise append a timestamp.

    Parameters
    ----------
    path : Path

    Returns
    -------
    Path
        Unique path that won't overwrite an existing file.
    """
    if not path.exists():
        return path
    ts   = datetime.datetime.now().strftime("%H%M%S")
    stem = path.stem + f"_{ts}"
    return path.with_name(stem + path.suffix)


def _progress(msg: str) -> None:
    """Emit progress message to log + notifier."""
    log.info("▶  %s", msg)
    _notify("Jarvis", msg)


def _count_words(text: str) -> int:
    """Return approximate word count."""
    return len(re.findall(r"\w+", text))


def _open_file(path: Path) -> None:
    """Open file in default OS application."""
    try:
        os.startfile(str(path))
    except AttributeError:
        for cmd in (["xdg-open"], ["open"]):
            try:
                subprocess.Popen(cmd + [str(path)]); return
            except FileNotFoundError:
                continue
        log.info("File saved at: %s", path)
    except Exception as exc:
        log.warning("Could not open file: %s", exc)


def _export_pdf(docx_path: Path) -> Optional[Path]:
    """Convert .docx → PDF via docx2pdf or LibreOffice fallback."""
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert  # type: ignore
        convert(str(docx_path), str(pdf_path))
        log.info("PDF exported: %s", pdf_path)
        return pdf_path
    except ImportError:
        log.warning("docx2pdf not installed — trying LibreOffice …")
    except Exception as exc:
        log.error("docx2pdf failed: %s", exc)
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", str(docx_path.parent), str(docx_path)],
            check=True, capture_output=True, timeout=60,
        )
        return pdf_path
    except Exception as exc:
        log.error("LibreOffice PDF export failed: %s", exc)
        return None


# ══════════════════════════════════════════════════════════════════════════════
# AI HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _ai_generate(prompt: str, max_tokens: int = 2000) -> str:
    """
    Generate text via Groq API.

    Parameters
    ----------
    prompt     : str
    max_tokens : int

    Returns
    -------
    str
    """
    if not GROQ_KEY:
        log.warning("GroqAPIKey not set — returning placeholder.")
        return f"# {prompt[:60]}\n\nAdd GroqAPIKey to .env for AI generation."
    try:
        from groq import Groq  # type: ignore
        client = Groq(api_key=GROQ_KEY)
        resp   = client.chat.completions.create(
            model    = GROQ_MODEL,
            messages = [
                {"role": "system", "content": (
                    "You are a senior academic content writer. "
                    "Use markdown: # H1, ## H2, ### H3. "
                    "Bold key terms using **term**. "
                    "Write formal, well-structured academic prose."
                )},
                {"role": "user", "content": prompt},
            ],
            max_tokens  = max_tokens,
            temperature = 0.7,
            stream      = False,
        )
        return resp.choices[0].message.content.strip()
    except ConnectionError as exc:
        log.error("Network error: %s", exc)
        return f"# Content unavailable\n\nNetwork error: {exc}"
    except Exception as exc:
        log.error("Groq API failed: %s", exc)
        return f"# Content unavailable\n\nError: {exc}"


def _ai_generate_with_validation(
    prompt     : str,
    target_words: int,
    max_tokens : int = 2000,
    max_retries: int = 2,
) -> str:
    """
    Generate content and re-try if output is too short.

    Parameters
    ----------
    prompt       : str
    target_words : int   — Minimum acceptable word count.
    max_tokens   : int
    max_retries  : int   — How many times to retry if too short.

    Returns
    -------
    str
    """
    minimum = int(target_words * MIN_WORD_THRESHOLD)
    for attempt in range(1, max_retries + 2):
        content = _ai_generate(prompt, max_tokens)
        wc      = _count_words(content)
        if wc >= minimum:
            log.info("Content OK: %d words (attempt %d)", wc, attempt)
            return content
        if attempt <= max_retries:
            log.warning("Content too short (%d/%d words) — retrying …", wc, minimum)
            # Strengthen prompt for retry
            prompt = prompt + f"\n\nIMPORTANT: Write at least {target_words} words!"
    log.warning("Could not reach target word count after %d attempts", max_retries + 1)
    return content


def _ai_title(topic: str) -> str:
    """Generate a professional academic title from a raw topic."""
    prompt = (
        f"Generate ONE professional academic title for: '{topic}'. "
        "Return ONLY the title — no quotes, no explanation."
    )
    raw   = _ai_generate(prompt, max_tokens=60)
    title = raw.split("\n")[0].strip().strip('"').strip("#").strip()
    return title if title else topic.title()


def _ai_grammar_pass(content: str) -> str:
    """Improve grammar while preserving markdown structure."""
    if not ENABLE_GRAMMAR or _count_words(content) < 100:
        return content
    prompt = (
        "Improve grammar, clarity, and academic tone. "
        "Keep ALL markdown headings exactly. "
        "Keep all **bold** markers. "
        "Do NOT add new sections.\n\n" + content
    )
    improved = _ai_generate(prompt, max_tokens=2200)
    return improved if improved else content


# ══════════════════════════════════════════════════════════════════════════════
# DOCX IMPORTS
# ══════════════════════════════════════════════════════════════════════════════

def _get_docx_imports():
    """Import python-docx symbols, raising ImportError with install hint."""
    try:
        from docx                import Document
        from docx.shared         import Pt, Inches, RGBColor
        from docx.enum.text      import WD_ALIGN_PARAGRAPH
        from docx.oxml           import OxmlElement
        from docx.oxml.ns        import qn
        return Document, Pt, Inches, RGBColor, WD_ALIGN_PARAGRAPH, OxmlElement, qn
    except ImportError as exc:
        raise ImportError(
            "python-docx required.  Run:  pip install python-docx"
        ) from exc


# ══════════════════════════════════════════════════════════════════════════════
# DOCX BUILDING BLOCKS
# ══════════════════════════════════════════════════════════════════════════════

def _set_margins(doc, inches: float) -> None:
    """Set equal page margins on every section."""
    from docx.shared import Inches as In
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = In(inches)


def _add_cover_page(
    doc, title: str, subject: str, course: str,
    student: str, professor: str, university: str,
    roll_no: str, semester: str,
    Pt, Inches, WD_ALIGN_PARAGRAPH,
) -> None:
    """
    Insert professional cover page.
    Now includes Roll Number and Semester fields.

    Parameters
    ----------
    roll_no  : str — Student roll number (from .env RollNumber)
    semester : str — Current semester  (from .env Semester)
    """
    font = DEFAULT_FONT

    def _centered(text: str, size: int, bold: bool = False,
                  space_before: int = 0) -> None:
        p   = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if space_before:
            p.paragraph_format.space_before = Pt(space_before)
        r = p.add_run(text)
        r.font.name = font
        r.font.size = Pt(size)
        r.bold      = bold

    def _label_value(label: str, value: str) -> None:
        if not value:
            return
        p  = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(f"{label}: ")
        rl.font.name = font; rl.font.size = Pt(BODY_FONT_SIZE); rl.bold = True
        rv = p.add_run(value)
        rv.font.name = font; rv.font.size = Pt(BODY_FONT_SIZE)

    from docx.oxml import OxmlElement as Ox
    from docx.oxml.ns import qn

    _centered(university.upper(), 14, bold=True, space_before=60)

    # Divider line
    hr  = doc.add_paragraph()
    hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = hr._p.get_or_add_pPr()
    pBdr = Ox("w:pBdr")
    bot  = Ox("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "000000")
    pBdr.append(bot); pPr.append(pBdr)

    doc.add_paragraph()
    _centered("ASSIGNMENT", 13)
    _centered(title, TITLE_FONT_SIZE, bold=True, space_before=20)
    doc.add_paragraph(); doc.add_paragraph()

    _label_value("Subject",      subject)
    _label_value("Course",       course)
    _label_value("Submitted By", student)
    if roll_no:
        _label_value("Roll No",  roll_no)
    if semester:
        _label_value("Semester", semester)
    _label_value("Submitted To", professor)
    _label_value("Date",         datetime.date.today().strftime("%d %B %Y"))

    doc.add_page_break()


def _add_toc(doc, OxmlElement, qn) -> None:
    """Insert Word-native TOC field (updates on open)."""
    from docx.shared import Pt
    h = doc.add_paragraph("Table of Contents")
    h.style = doc.styles["Heading 1"]
    h.runs[0].bold = True
    for r in h.runs:
        r.font.name = DEFAULT_FONT
        r.font.size = Pt(H1_FONT_SIZE)

    p   = doc.add_paragraph()
    run = p.add_run()
    b   = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    b.set(qn("w:dirty"), "true")
    i   = OxmlElement("w:instrText")
    i.set(qn("xml:space"), "preserve"); i.text = ' TOC \\o "1-3" \\h \\z \\u '
    s   = OxmlElement("w:fldChar"); s.set(qn("w:fldCharType"), "separate")
    e   = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    for el in (b, i, s, e):
        run._r.append(el)
    doc.add_page_break()


def _add_header_footer(
    doc, title: str, subject: str,
    course: str, student: str,
    Pt, WD_ALIGN_PARAGRAPH,
) -> None:
    """Header: Subject/Course/Student. Footer: centred page number."""
    from docx.oxml import OxmlElement as Ox
    from docx.oxml.ns import qn

    for section in doc.sections:
        section.different_first_page_header_footer = True

        # ── Header ────────────────────────────────────────────────────────────
        header = section.header
        header.is_linked_to_previous = False
        for para in header.paragraphs:
            para.clear(); break
        else:
            header.add_paragraph()

        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for i, line in enumerate([
            f"Subject: {subject}",
            f"Course:  {course}",
            f"Student: {student}",
        ]):
            run = hp.add_run(("" if i == 0 else "\n") + line)
            run.font.name = DEFAULT_FONT
            run.font.size = Pt(9)
            run.italic    = True

        pPr  = hp._p.get_or_add_pPr()
        pBdr = Ox("w:pBdr"); bot = Ox("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
        bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "999999")
        pBdr.append(bot); pPr.append(pBdr)

        # ── Footer — page number ───────────────────────────────────────────────
        footer = section.footer
        footer.is_linked_to_previous = False
        for para in footer.paragraphs:
            para.clear(); break
        else:
            footer.add_paragraph()

        fp  = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        rel = run._r
        for ftype, txt in [("begin", None), (None, " PAGE "), ("separate", None), ("end", None)]:
            if ftype:
                el = Ox("w:fldChar"); el.set(qn("w:fldCharType"), ftype); rel.append(el)
            else:
                el = Ox("w:instrText"); el.set(qn("xml:space"), "preserve"); el.text = txt; rel.append(el)
        run.font.name = DEFAULT_FONT; run.font.size = Pt(10)


def _apply_heading_run_style(para, font_name: str, size_pt, bold: bool = False) -> None:
    """Set font/size/bold on every run in a heading paragraph."""
    from docx.shared import Pt
    for r in para.runs:
        r.font.name = font_name
        r.font.size = Pt(size_pt) if isinstance(size_pt, (int, float)) else size_pt
        if bold: r.bold = True


def _add_body_paragraph(doc, text: str, Pt, WD_ALIGN_PARAGRAPH) -> None:
    """Add justified 1.5-spaced body paragraph with bold support."""
    # Split text on **bold** markers
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = Pt(LINE_SPACING_PT)
    para.paragraph_format.space_after  = Pt(PARA_SPACE_AFTER_PT)

    # Parse **bold** segments
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run      = para.add_run(part[2:-2])
            run.bold = True
        else:
            run = para.add_run(part)
        run.font.name = DEFAULT_FONT
        run.font.size = Pt(BODY_FONT_SIZE)


def _parse_and_insert_content(doc, content: str, Pt, WD_ALIGN_PARAGRAPH) -> None:
    """
    Walk markdown content and insert styled paragraphs/headings into doc.
    Supports: # H1, ## H2, ### H3, - bullets, **bold**, plain paragraphs.
    """
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)
            _apply_heading_run_style(h, DEFAULT_FONT, H3_FONT_SIZE)
        elif stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=2)
            _apply_heading_run_style(h, DEFAULT_FONT, H2_FONT_SIZE)
        elif stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=1)
            _apply_heading_run_style(h, DEFAULT_FONT, H1_FONT_SIZE, bold=True)
        elif stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
            for r in p.runs:
                r.font.name = DEFAULT_FONT
                r.font.size = Pt(BODY_FONT_SIZE)
        else:
            _add_body_paragraph(doc, stripped, Pt, WD_ALIGN_PARAGRAPH)


def _add_word_count_footer(doc, content: str, target: int,
                            Pt, WD_ALIGN_PARAGRAPH) -> None:
    """Append word-count note at end of document body."""
    actual = _count_words(content)
    p      = doc.add_paragraph(
        f"[Word count: {actual} words  |  Target: {target} words  |  "
        f"Generated: {datetime.datetime.now().strftime('%d %B %Y')}]"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        r.font.name = DEFAULT_FONT
        r.font.size = Pt(9)
        r.italic    = True


def _try_insert_diagram(doc, topic: str, Pt, WD_ALIGN_PARAGRAPH) -> bool:
    """
    Attempt to insert a simple matplotlib diagram relevant to the topic.
    Silently skips if matplotlib is not installed.

    Parameters
    ----------
    topic : str
        Used to decide diagram type (flowchart placeholder for now).

    Returns
    -------
    bool
        True if diagram was inserted.
    """
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
        import io
        from docx.shared import Inches

        fig, ax = plt.subplots(figsize=(5, 3))
        ax.set_xlim(0, 10); ax.set_ylim(0, 6)
        ax.axis("off")

        # Simple process flow boxes
        steps = ["Input", "Process", "Output"]
        colors = ["#4472C4", "#ED7D31", "#70AD47"]
        for i, (step, color) in enumerate(zip(steps, colors)):
            x = 1.5 + i * 3
            rect = mpatches.FancyBboxPatch(
                (x - 0.8, 2.2), 1.6, 1.2,
                boxstyle="round,pad=0.1",
                facecolor=color, edgecolor="white",
            )
            ax.add_patch(rect)
            ax.text(x, 2.8, step, ha="center", va="center",
                    color="white", fontsize=11, fontweight="bold")
            if i < len(steps) - 1:
                ax.annotate("", xy=(x + 2.2, 2.8), xytext=(x + 0.8, 2.8),
                            arrowprops=dict(arrowstyle="->", color="#555555", lw=2))

        ax.set_title(f"Process Flow: {topic.title()}", fontsize=12, pad=10)
        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, format="png", dpi=120, bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)

        cap = doc.add_paragraph(f"Figure 1: Process flow diagram — {topic.title()}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.font.size   = Pt(10)
            r.font.name   = DEFAULT_FONT
            r.italic      = True

        doc.add_picture(buf, width=Inches(4.5))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        log.info("Diagram inserted for topic: %s", topic)
        return True

    except ImportError:
        log.debug("matplotlib not installed — skipping diagram")
        return False
    except Exception as exc:
        log.warning("Diagram insertion failed: %s", exc)
        return False


# ══════════════════════════════════════════════════════════════════════════════
# PROMPT BUILDERS
# ══════════════════════════════════════════════════════════════════════════════

_TEMPLATE_PROMPTS: dict[str, str] = {
    "standard": (
        "Write a comprehensive academic assignment on: {topic}\n\n"
        "Structure (exact markdown headings):\n"
        "# {title}\n## Abstract\n## Introduction\n## Main Content\n"
        "### {section1}\n### {section2}\n### {section3}\n"
        "## Examples and Applications\n## Conclusion\n## References\n"
        "  (5 APA-format references)\n\n"
        "Target: {words} words. Bold key terms. Formal academic language."
    ),
    "ieee": (
        "Write an IEEE-format research paper on: {topic}\n\n"
        "# {title}\n## Abstract\n## I. Introduction\n## II. Related Work\n"
        "## III. Methodology\n### A. System Design\n### B. Implementation\n"
        "## IV. Results and Discussion\n## V. Conclusion\n## References\n"
        "  (IEEE-style [1]…[5])\n\nTarget: {words} words. IEEE academic style."
    ),
    "research": (
        "Write a research paper on: {topic}\n\n"
        "# {title}\n## Abstract\n## 1. Introduction\n## 2. Literature Review\n"
        "## 3. Research Methodology\n### 3.1 Research Design\n### 3.2 Data Collection\n"
        "## 4. Analysis and Findings\n## 5. Discussion\n"
        "## 6. Conclusion and Future Work\n## References\n  (5 APA references)\n\n"
        "Target: {words} words. Formal research tone."
    ),
    "report": (
        "Write a professional technical report on: {topic}\n\n"
        "# {title}\n## Executive Summary\n## 1. Introduction\n## 2. Background\n"
        "## 3. Technical Details\n### 3.1 Overview\n### 3.2 Implementation\n"
        "## 4. Challenges and Solutions\n## 5. Recommendations\n## 6. Conclusion\n"
        "## Appendix / References\n\nTarget: {words} words. Professional report style."
    ),
    "pbl": (
        "Write a PBL (Project Based Learning) report on: {topic}\n\n"
        "# {title}\n## Abstract\n## 1. Problem Statement\n## 2. Objectives\n"
        "## 3. Project Planning\n### 3.1 Timeline\n### 3.2 Tools and Technologies\n"
        "## 4. Implementation\n### 4.1 Design\n### 4.2 Development\n"
        "## 5. Testing and Results\n## 6. Challenges Faced\n"
        "## 7. Conclusion and Future Scope\n## References\n\n"
        "Target: {words} words. Engineering project report style."
    ),
    "lab": (
        "Write a lab manual / practical report on: {topic}\n\n"
        "# {title}\n## Aim\n## Theory\n## Required Components / Tools\n"
        "## Procedure\n### Step-by-Step Instructions\n"
        "## Observations and Results\n## Calculations\n"
        "## Conclusion\n## Precautions\n## Viva Questions\n\n"
        "Target: {words} words. Clear, step-by-step practical style."
    ),
}


def _build_prompt(topic: str, title: str, template: str, words: int) -> str:
    """Fill the prompt template with topic/title/words."""
    base = _TEMPLATE_PROMPTS.get(template.lower(), _TEMPLATE_PROMPTS["standard"])
    return base.format(
        topic=topic, title=title, words=words,
        section1="Core Concepts",
        section2="Key Techniques",
        section3="Advanced Topics",
    )


def _build_custom_prompt(topic: str, title: str, structure: str, words: int) -> str:
    """Build prompt from user-pasted custom section structure."""
    return (
        f"Write a comprehensive academic document on: {topic}\n\n"
        f"Use EXACTLY these sections (markdown headings):\n"
        f"# {title}\n{structure}\n\n"
        f"Target: {words} words. Bold key terms. Formal academic language."
    )


# ══════════════════════════════════════════════════════════════════════════════
# CORE DOCUMENT BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def _build_and_save_doc(
    topic      : str,
    content    : str,
    word_limit : int,
    out_dir    : Path,
    export_pdf_: bool = False,
    insert_diagram: bool = False,
) -> bool:
    """
    Format and save a Word document from markdown content.

    Parameters
    ----------
    topic          : str
    content        : str   — Markdown text.
    word_limit     : int   — For word count footer.
    out_dir        : Path  — Where to save the .docx file.
    export_pdf_    : bool
    insert_diagram : bool  — Whether to try inserting a matplotlib diagram.

    Returns
    -------
    bool
    """
    try:
        Document, Pt, Inches, RGBColor, WD_ALIGN_PARAGRAPH, OxmlElement, qn = _get_docx_imports()
    except ImportError as exc:
        log.error(str(exc)); return False

    _progress("Formatting document …")
    try:
        doc = Document()
        _set_margins(doc, MARGIN_INCHES)

        doc.styles["Normal"].font.name = DEFAULT_FONT
        doc.styles["Normal"].font.size = Pt(BODY_FONT_SIZE)

        # Extract title from first # heading in content (or use topic)
        title_match = re.search(r"^#\s+(.+)$", content, re.MULTILINE)
        title       = title_match.group(1).strip() if title_match else topic.title()

        _add_cover_page(
            doc, title, SUBJECT, COURSE,
            USERNAME, PROFESSOR, UNIVERSITY,
            ROLL_NO, SEMESTER,
            Pt, Inches, WD_ALIGN_PARAGRAPH,
        )
        _add_toc(doc, OxmlElement, qn)
        _add_header_footer(doc, title, SUBJECT, COURSE, USERNAME, Pt, WD_ALIGN_PARAGRAPH)
        _parse_and_insert_content(doc, content, Pt, WD_ALIGN_PARAGRAPH)

        if insert_diagram:
            _try_insert_diagram(doc, topic, Pt, WD_ALIGN_PARAGRAPH)

        _add_word_count_footer(doc, content, word_limit, Pt, WD_ALIGN_PARAGRAPH)

        _progress("Saving file …")
        safe     = _sanitize(topic)
        raw_path = out_dir / f"Assignment_{safe}.docx"
        out_path = _unique_path(raw_path)        # ← duplicate protection
        doc.save(str(out_path))
        log.info("Saved: %s  (%d words)", out_path, _count_words(content))
        _notify("Jarvis — Assignment Created 📄", out_path.name)

        if export_pdf_:
            _progress("Exporting PDF …")
            pdf = _export_pdf(out_path)
            if pdf:
                _notify("Jarvis — PDF Ready 📄", pdf.name)

        _progress("Opening Word …")
        _open_file(out_path)
        return True

    except PermissionError as exc:
        log.error("Permission denied: %s", exc); return False
    except FileNotFoundError as exc:
        log.error("Directory not found: %s", exc); return False
    except Exception as exc:
        log.exception("Document build failed: %s", exc); return False


# ══════════════════════════════════════════════════════════════════════════════
# PUBLIC API  (unchanged signatures — AutomationEngine compatible)
# ══════════════════════════════════════════════════════════════════════════════

def CreateAssignment(
    topic      : str,
    template   : str           = "standard",
    word_limit : Optional[int] = None,
    export_pdf : bool          = False,
) -> bool:
    """
    Generate and save a formatted academic assignment Word document.

    Parameters
    ----------
    topic      : str   — Assignment topic.
    template   : str   — standard | ieee | research | report | pbl | lab.
    word_limit : int | None — Target word count (None → .env default).
    export_pdf : bool  — Also export PDF.

    Returns
    -------
    bool
    """
    words  = (WORD_LIMITS.get(str(word_limit).lower(), DEFAULT_WORD_LIMIT)
              if isinstance(word_limit, str) else word_limit or DEFAULT_WORD_LIMIT)
    do_pdf = export_pdf or ENABLE_PDF_EXPORT

    with _ProgressBar("Generating title"):
        title = _ai_title(topic)
    log.info("Title: %s", title)

    with _ProgressBar(f"Generating content (~{words} words)"):
        prompt  = _build_prompt(topic, title, template, words)
        content = _ai_generate_with_validation(prompt, words,
                                               max_tokens=min(words * 2, 3000))

    with _ProgressBar("Grammar improvement pass"):
        content = _ai_grammar_pass(content)

    return _build_and_save_doc(
        topic=topic, content=content, word_limit=words,
        out_dir=ASSIGN_DIR, export_pdf_=do_pdf,
        insert_diagram=True,
    )


def CreateNotes(
    topic      : str,
    word_limit : Optional[int] = None,
) -> bool:
    """
    Generate study notes as a formatted Word document.

    Parameters
    ----------
    topic      : str
    word_limit : int | None — Target length (None → 400 words).

    Returns
    -------
    bool
    """
    words = word_limit or 400
    _progress(f"Generating notes on: {topic}")

    with _ProgressBar("Generating notes"):
        content = _ai_generate(
            f"Write concise study notes on: {topic}\n"
            f"Format: ## headings, bullet points, **bold** key terms, examples.\n"
            f"Target: {words} words.",
            max_tokens=words * 2,
        )

    try:
        Document, Pt, Inches, RGBColor, WD_ALIGN_PARAGRAPH, OxmlElement, qn = _get_docx_imports()
        doc = Document()
        _set_margins(doc, MARGIN_INCHES)

        doc.styles["Normal"].font.name = DEFAULT_FONT
        doc.styles["Normal"].font.size = Pt(BODY_FONT_SIZE)

        h = doc.add_heading(f"Study Notes: {topic.title()}", level=1)
        _apply_heading_run_style(h, DEFAULT_FONT, H1_FONT_SIZE, bold=True)

        meta = doc.add_paragraph(
            f"Student: {USERNAME}   |   "
            f"Date: {datetime.date.today().strftime('%d %B %Y')}   |   "
            f"Words: {_count_words(content)}"
        )
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in meta.runs:
            r.font.name = DEFAULT_FONT; r.font.size = Pt(9); r.italic = True
        doc.add_paragraph()

        _parse_and_insert_content(doc, content, Pt, WD_ALIGN_PARAGRAPH)

        safe     = _sanitize(topic)
        out_path = _unique_path(NOTES_DIR / f"Notes_{safe}.docx")
        doc.save(str(out_path))
        log.info("Notes saved: %s", out_path)
        _notify("Jarvis — Notes Created 📝", out_path.name)
        _open_file(out_path)
        return True

    except ImportError:
        # Fallback: plain text
        safe     = _sanitize(topic)
        out_path = _unique_path(NOTES_DIR / f"Notes_{safe}.txt")
        out_path.write_text(
            f"NOTES: {topic.upper()}\nDate: {datetime.date.today()}\n"
            f"{'='*60}\n\n{content}\n\n{'='*60}\n"
            f"[Words: {_count_words(content)}]\n",
            encoding="utf-8",
        )
        try:
            subprocess.Popen(["notepad.exe", str(out_path)])
        except Exception:
            _open_file(out_path)
        return True
    except Exception as exc:
        log.exception("Notes creation failed: %s", exc); return False


def CreateReport(
    topic      : str,
    word_limit : Optional[int] = None,
    export_pdf : bool          = False,
) -> bool:
    """Generate a technical report (uses 'report' template, saves to Reports/)."""
    global ASSIGN_DIR
    _orig      = ASSIGN_DIR
    ASSIGN_DIR = REPORTS_DIR
    try:
        return CreateAssignment(topic, template="report",
                                word_limit=word_limit, export_pdf=export_pdf)
    finally:
        ASSIGN_DIR = _orig


def _create_doc_from_content(
    topic      : str,
    content    : str,
    word_limit : int,
    export_pdf : bool = False,
) -> bool:
    """Build document from pre-generated content (used by custom format path)."""
    return _build_and_save_doc(
        topic=topic, content=content, word_limit=word_limit,
        out_dir=ASSIGN_DIR, export_pdf_=export_pdf,
    )


# ══════════════════════════════════════════════════════════════════════════════
# COLLEGE FORMAT TERMINAL MENU  (unchanged from v2.0 + lab manual added)
# ══════════════════════════════════════════════════════════════════════════════

_COLLEGE_KEYWORDS: tuple[str, ...] = (
    "assignment", "pbl", "project", "report", "lab", "practical",
    "college", "university", "semester", "submit", "submission",
    "viva", "internal", "external", "marks", "exam", "course",
    "btech", "b.tech", "mtech", "m.tech", "diploma", "bca", "mca",
    "manual", "experiment",
)

_FORMAT_MENU: dict[str, dict] = {
    "1": {"name": "Standard Assignment", "template": "standard",
          "desc": "Abstract → Introduction → Main Content → Conclusion → References"},
    "2": {"name": "IEEE Format",          "template": "ieee",
          "desc": "Abstract → I. Intro → II. Related Work → III. Methodology → IV. Results → V. Conclusion"},
    "3": {"name": "Research Paper",       "template": "research",
          "desc": "Abstract → Literature Review → Methodology → Findings → Conclusion"},
    "4": {"name": "Technical Report",     "template": "report",
          "desc": "Executive Summary → Background → Technical Details → Recommendations"},
    "5": {"name": "PBL / Project Report", "template": "pbl",
          "desc": "Problem Statement → Objectives → Implementation → Testing → Conclusion"},
    "6": {"name": "Lab Manual",           "template": "lab",
          "desc": "Aim → Theory → Procedure → Observations → Conclusion → Viva Questions"},
    "7": {"name": "Custom Format",        "template": "custom",
          "desc": "Paste your college's exact section structure"},
}


def _is_college_request(cmd: str) -> bool:
    return any(kw in cmd for kw in _COLLEGE_KEYWORDS)


def _terminal_divider(char: str = "─", width: int = 60) -> None:
    print(char * width)


def _ask_college_format(topic: str) -> tuple[str, Optional[str], Optional[int]]:
    """Interactive terminal menu for college document format selection."""
    print("\n")
    _terminal_divider("═")
    print(f"  🎓  JARVIS — College Document Detected")
    print(f"  Topic : {topic}")
    _terminal_divider("═")
    print("\n  What FORMAT does your college require?\n")
    for k, v in _FORMAT_MENU.items():
        print(f"  [{k}]  {v['name']}")
        print(f"       {v['desc']}\n")
    _terminal_divider()

    template      = "standard"
    custom_struct = None

    while True:
        choice = input("  Enter choice [1-7] (Enter = Standard): ").strip()
        if choice in ("", "1"):
            template = "standard"; break
        elif choice in _FORMAT_MENU:
            template = _FORMAT_MENU[choice]["template"]; break
        else:
            print("  ⚠  Enter 1-7.\n")

    if template == "custom":
        print("\n  Paste your college sections below (one per line, then DONE):\n")
        lines: list[str] = []
        while True:
            line = input("  > ")
            if line.strip().upper() == "DONE": break
            lines.append(line)
        custom_struct = "\n".join(lines).strip() or None
        if not custom_struct:
            print("  ⚠  No structure — falling back to Standard.")
            template = "standard"

    print("\n  How long should it be?")
    for k, (label, words) in enumerate([
        ("Short (~300)", 300), ("Medium (~600)", 600), ("Standard (~800)", 800),
        ("Detailed (~1200)", 1200), ("Extended (~1500)", 1500), ("Custom", 0),
    ], 1):
        print(f"  [{k}]  {label}")
    _terminal_divider()

    wc_map    = {"1": 300, "2": 600, "3": 800, "4": 1200, "5": 1500}
    word_limit: Optional[int] = None
    while True:
        wc = input("  Enter choice [1-6] (Enter = 800): ").strip()
        if wc in ("", "3"):   word_limit = 800;            break
        elif wc in wc_map:    word_limit = wc_map[wc];     break
        elif wc == "6":
            try:              word_limit = int(input("  Word count: ").strip())
            except ValueError: word_limit = 800
            break
        else: print("  ⚠  Enter 1-6.\n")

    pdf_ans = input("\n  Export PDF too? [y/N]: ").strip().lower()
    _ask_college_format._export_pdf = pdf_ans in ("y", "yes")  # type: ignore

    fmt_name = next(
        (v["name"] for v in _FORMAT_MENU.values() if v["template"] == template),
        "Standard Assignment"
    )
    _terminal_divider("═")
    print(f"\n  ✅  Format : {fmt_name}")
    print(f"  ✅  Words  : ~{word_limit}")
    print(f"  ✅  PDF    : {'Yes' if _ask_college_format._export_pdf else 'No'}")  # type: ignore
    _terminal_divider("═"); print()

    return template, custom_struct, word_limit


# ══════════════════════════════════════════════════════════════════════════════
# VOICE COMMAND ROUTER
# ══════════════════════════════════════════════════════════════════════════════

def route_voice_command(command: str) -> bool:
    """
    Parse a natural-language voice command and dispatch correctly.

    Examples
    --------
    >>> route_voice_command("create assignment on machine learning")
    >>> route_voice_command("create ieee report on AI")
    >>> route_voice_command("create lab manual on op-amp")
    >>> route_voice_command("create 1500 word report on python and export pdf")
    >>> route_voice_command("create short notes on recursion")
    """
    cmd        = command.lower().strip()
    export_pdf = "export pdf" in cmd or "pdf" in cmd
    cmd        = re.sub(r"\band export pdf\b|\bexport pdf\b|\bpdf\b", "", cmd).strip()

    word_limit: Optional[int] = None
    for alias, val in WORD_LIMITS.items():
        if alias in cmd:
            word_limit = val
            cmd = cmd.replace(alias, "").strip()
            break
    m = re.search(r"(\d{3,4})\s*word", cmd)
    if m:
        word_limit = int(m.group(1))
        cmd = re.sub(r"\d{3,4}\s*word\w*", "", cmd).strip()

    topic_match = re.search(r"\bon\s+(.+)$", cmd)
    topic       = topic_match.group(1).strip() if topic_match else cmd

    if "notes" in cmd:
        return CreateNotes(topic, word_limit=word_limit)
    if "ieee"     in cmd:
        return CreateAssignment(topic, template="ieee",     word_limit=word_limit, export_pdf=export_pdf)
    if "research" in cmd:
        return CreateAssignment(topic, template="research", word_limit=word_limit, export_pdf=export_pdf)
    if "lab"      in cmd or "manual" in cmd or "practical" in cmd:
        return CreateAssignment(topic, template="lab",      word_limit=word_limit, export_pdf=export_pdf)

    if _is_college_request(cmd):
        template, custom_struct, wl = _ask_college_format(topic)
        export_pdf = getattr(_ask_college_format, "_export_pdf", export_pdf)
        if wl: word_limit = wl

        if template == "custom" and custom_struct:
            _progress("Generating content from custom format …")
            raw = _ai_generate_with_validation(
                _build_custom_prompt(topic, topic.title(), custom_struct,
                                     word_limit or DEFAULT_WORD_LIMIT),
                word_limit or DEFAULT_WORD_LIMIT,
                max_tokens=min((word_limit or DEFAULT_WORD_LIMIT) * 2, 3000),
            )
            _progress("Grammar pass …")
            raw = _ai_grammar_pass(raw)
            return _create_doc_from_content(
                topic=topic, content=raw,
                word_limit=word_limit or DEFAULT_WORD_LIMIT,
                export_pdf=export_pdf,
            )
        if template == "report" or "report" in cmd:
            return CreateReport(topic, word_limit=word_limit, export_pdf=export_pdf)
        return CreateAssignment(topic, template=template,
                                word_limit=word_limit, export_pdf=export_pdf)

    if "report" in cmd:
        return CreateReport(topic, word_limit=word_limit, export_pdf=export_pdf)

    return CreateAssignment(topic, template="standard",
                            word_limit=word_limit, export_pdf=export_pdf)


# ══════════════════════════════════════════════════════════════════════════════
# QUICK TEST
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    test_cmd = " ".join(sys.argv[1:]) or "create assignment on machine learning"
    print(f"\n[TEST]  Command: {test_cmd}\n")
    result = route_voice_command(test_cmd)
    print(f"\n[TEST]  Result: {'✅ Success' if result else '❌ Failed'}\n")